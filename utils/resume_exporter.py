"""
Resume export functionality
Generates PDF, DOCX, and CSV exports
"""

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
import io
import re

def parse_resume_sections(resume_text: str) -> dict:
    """
    Parse resume text into sections

    Returns:
        {"name": str, "contact": str, "summary": str, "experience": str,
         "education": str, "skills": str, "certifications": str}
    """
    sections = {
        "name": "",
        "contact": "",
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "certifications": ""
    }

    lines = resume_text.split('\n')

    # Extract name (usually first line)
    if lines:
        sections["name"] = lines[0].strip()[:100]

    # Find sections by keywords
    current_section = None
    section_content = []

    for line in lines[1:]:
        line_lower = line.lower()

        # Detect section headers
        if re.search(r'\b(contact|email|phone|linkedin|location)\b', line_lower) and not current_section:
            current_section = "contact"
        elif re.search(r'\b(summary|objective|profile)\b', line_lower):
            if current_section and section_content:
                sections[current_section] = '\n'.join(section_content)
            current_section = "summary"
            section_content = []
        elif re.search(r'\b(experience|employment|work history|work)\b', line_lower):
            if current_section and section_content:
                sections[current_section] = '\n'.join(section_content)
            current_section = "experience"
            section_content = []
        elif re.search(r'\b(education|academic|degree|university|school)\b', line_lower):
            if current_section and section_content:
                sections[current_section] = '\n'.join(section_content)
            current_section = "education"
            section_content = []
        elif re.search(r'\b(skills|technical|core|abilities)\b', line_lower):
            if current_section and section_content:
                sections[current_section] = '\n'.join(section_content)
            current_section = "skills"
            section_content = []
        elif re.search(r'\b(certification|certificate|credential|license)\b', line_lower):
            if current_section and section_content:
                sections[current_section] = '\n'.join(section_content)
            current_section = "certifications"
            section_content = []
        else:
            if current_section:
                section_content.append(line)

    # Store last section
    if current_section and section_content:
        sections[current_section] = '\n'.join(section_content)

    return sections

def generate_pdf_resume(resume_text: str, template: str = "modern") -> bytes:
    """
    Generate PDF resume

    Args:
        resume_text: Resume text content
        template: "modern", "classic", or "minimal"

    Returns:
        PDF bytes
    """
    # Parse sections
    sections = parse_resume_sections(resume_text)

    # Create PDF in memory
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter,
                           rightMargin=0.5*inch, leftMargin=0.5*inch,
                           topMargin=0.5*inch, bottomMargin=0.5*inch)

    elements = []

    # Define styles
    if template == "modern":
        name_style = ParagraphStyle(
            'CustomName',
            fontSize=24,
            textColor=colors.HexColor('#3B82F6'),
            spaceAfter=6,
            alignment=1  # Center
        )
        section_style = ParagraphStyle(
            'CustomSection',
            fontSize=13,
            textColor=colors.HexColor('#111827'),
            spaceAfter=8,
            spaceBefore=10,
            borderColor=colors.HexColor('#3B82F6'),
            borderBottomWidth=2,
            borderBottomPadding=3
        )
    elif template == "classic":
        name_style = ParagraphStyle(
            'CustomName',
            fontSize=20,
            textColor=colors.black,
            spaceAfter=4,
            alignment=1
        )
        section_style = ParagraphStyle(
            'CustomSection',
            fontSize=12,
            textColor=colors.black,
            spaceAfter=6,
            spaceBefore=8,
            borderBottomWidth=1,
            borderBottomColor=colors.black
        )
    else:  # minimal
        name_style = ParagraphStyle(
            'CustomName',
            fontSize=18,
            textColor=colors.black,
            spaceAfter=3,
            alignment=1
        )
        section_style = ParagraphStyle(
            'CustomSection',
            fontSize=11,
            textColor=colors.black,
            spaceAfter=6,
            spaceBefore=6
        )

    # Add name
    if sections["name"]:
        elements.append(Paragraph(sections["name"], name_style))
        elements.append(Spacer(1, 0.1*inch))

    # Add contact
    if sections["contact"]:
        contact_style = ParagraphStyle('Contact', fontSize=9, textColor=colors.grey)
        elements.append(Paragraph(sections["contact"][:200], contact_style))
        elements.append(Spacer(1, 0.15*inch))

    # Add sections
    section_titles = ["summary", "experience", "education", "skills", "certifications"]
    section_labels = ["SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS", "CERTIFICATIONS"]

    body_style = ParagraphStyle('Body', fontSize=10, spaceAfter=4, leading=12)

    for label, key in zip(section_labels, section_titles):
        if sections[key]:
            elements.append(Paragraph(label, section_style))

            content = sections[key][:1000]  # Limit to 1000 chars per section
            elements.append(Paragraph(content, body_style))
            elements.append(Spacer(1, 0.1*inch))

    # Build PDF
    doc.build(elements)
    pdf_buffer.seek(0)

    return pdf_buffer.getvalue()

def generate_docx_resume(resume_text: str) -> bytes:
    """
    Generate DOCX resume

    Args:
        resume_text: Resume text content

    Returns:
        DOCX bytes
    """
    # Parse sections
    sections = parse_resume_sections(resume_text)

    # Create DOCX in memory
    doc = Document()

    # Add name
    if sections["name"]:
        name_para = doc.add_paragraph(sections["name"])
        name_para.style = 'Heading 1'
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add contact
    if sections["contact"]:
        contact_para = doc.add_paragraph(sections["contact"][:200])
        contact_para.style = 'Normal'
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Add sections
    section_titles = ["summary", "experience", "education", "skills", "certifications"]
    section_labels = ["Summary", "Experience", "Education", "Skills", "Certifications"]

    for label, key in zip(section_labels, section_titles):
        if sections[key]:
            # Add section header
            header = doc.add_paragraph(label)
            header.style = 'Heading 2'
            header_run = header.runs[0]
            header_run.font.color.rgb = RGBColor(59, 130, 246)

            # Add content
            content = sections[key][:1000]
            content_para = doc.add_paragraph(content)
            content_para.style = 'Normal'

            doc.add_paragraph()  # Spacing

    # Get bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer.getvalue()

def generate_csv_export(analysis_data: dict) -> str:
    """
    Generate CSV export of analysis

    Args:
        analysis_data: Analysis results dict

    Returns:
        CSV string
    """
    if not analysis_data:
        return ""

    output = io.StringIO()
    writer = csv.writer(output)

    # Write header
    writer.writerow(["Field", "Value"])

    # Write data
    for key, value in analysis_data.items():
        if isinstance(value, (list, dict)):
            value = str(value)[:200]  # Truncate

        writer.writerow([key, str(value)[:500]])

    return output.getvalue()
